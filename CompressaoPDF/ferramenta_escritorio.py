import os
import shutil
import subprocess
import tempfile
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from datetime import datetime
import webbrowser
import pyautogui
import pyperclip
import time
import re
from docx import Document
from docx2pdf import convert
from PyPDF2 import PdfReader, PdfWriter
from docx.oxml.ns import qn

# ========== Funções PDF ==========
def unir_pdfs(lista_arquivos, pdf_saida):
    pdf_writer = PdfWriter()
    arquivos_pdf = [f for f in lista_arquivos if f.lower().endswith(".pdf")]
    if not pdf_saida:
        messagebox.showerror("Erro interno", "Caminho de saída inválido.")
        return False

    if not arquivos_pdf:
        messagebox.showerror("Erro", "Nenhum PDF selecionado.")
        return False

    for caminho_pdf in arquivos_pdf:
        try:
            leitor = PdfReader(caminho_pdf)
            for pagina in leitor.pages:
                pdf_writer.add_page(pagina)
        except Exception as e:
            print(f"Erro ao ler {caminho_pdf}: {e}")

    try:
        with open(pdf_saida, 'wb') as saida:
            pdf_writer.write(saida)
        return True
    except Exception as e:
        messagebox.showerror("Erro ao salvar PDF unido", str(e))
        return False

def reduzir_pdf(input_pdf, output_pdf):
    qualidade = "/ebook"
    pasta_temp = tempfile.mkdtemp(prefix="temp_gs_")
    nome_temp = os.path.basename(output_pdf).replace(".pdf", "_temp.pdf")
    temp_output = os.path.join(pasta_temp, nome_temp)

    gs_path = None
    possiveis_caminhos = [
        r"C:\\Program Files\\gs\\gs10.05.1\\bin\\gswin64c.exe",
        r"C:\\Program Files (x86)\\gs\\gs10.05.1\\bin\\gswin32c.exe"
    ]
    for caminho in possiveis_caminhos:
        if os.path.exists(caminho):
            gs_path = caminho
            break

    if not gs_path:
        messagebox.showerror("Erro", "Ghostscript não encontrado.")
        return

    comandos = [
        f'"{gs_path}"',
        "-sDEVICE=pdfwrite",
        "-dCompatibilityLevel=1.4",
        f"-dPDFSETTINGS={qualidade}",
        "-dNOPAUSE", "-dBATCH", "-dQUIET",
        "-dDownsampleColorImages=true",
        "-dColorImageDownsampleType=/Bicubic",
        "-dColorImageResolution=110",
        "-dDownsampleGrayImages=true",
        "-dGrayImageDownsampleType=/Bicubic",
        "-dGrayImageResolution=110",
        "-dDownsampleMonoImages=true",
        "-dMonoImageDownsampleType=/Subsample",
        "-dMonoImageResolution=110",
        f'-sOutputFile="{temp_output}"',
        f'"{input_pdf}"'
    ]

    try:
        subprocess.run(" ".join(comandos), shell=True, check=True)
        if not os.path.isfile(temp_output):
            messagebox.showerror("Erro", "Arquivo comprimido não foi gerado.")
            return
        os.replace(temp_output, output_pdf)
        tamanho_mb = os.path.getsize(output_pdf) / (1024 * 1024)
        messagebox.showinfo("Sucesso", f"PDF comprimido salvo com {tamanho_mb:.2f} MB\n\n{output_pdf}")
    except subprocess.CalledProcessError as e:
        messagebox.showerror("Erro Ghostscript", str(e))
    finally:
        shutil.rmtree(pasta_temp, ignore_errors=True)

# ========== Função auxiliar para executar união e compressão ==========
def executar_pdf_tool():
    arquivos = filedialog.askopenfilenames(title="Selecione os PDFs", filetypes=[["PDFs", "*.pdf"]])
    if not arquivos:
        return
    destino = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[["PDF", "*.pdf"]], title="Salvar como")
    if not destino:
        messagebox.showerror("Erro", "Nenhum destino selecionado para salvar o PDF.")
        return

    temp_pdf = os.path.join(tempfile.gettempdir(), "pdf_temp_unido.pdf")
    if unir_pdfs(arquivos, temp_pdf):
        reduzir_pdf(temp_pdf, destino)

# ========== Funções auxiliares ==========
def gerar_data_formatada():
    meses = ['janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
             'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro']
    hoje = datetime.now()
    return f"{hoje.day} de {meses[hoje.month - 1]} de {hoje.year}"

def inserir_texto_com_estilo(paragrafo, texto_total):
    partes = re.split(r"(\|\|NEGRITO_INICIO\|\|.*?\|\|NEGRITO_FIM\|\||\|\|ITALICO_INICIO\|\|.*?\|\|ITALICO_FIM\|\|)", texto_total)
    for parte in partes:
        if parte.startswith("||NEGRITO_INICIO||") and parte.endswith("||NEGRITO_FIM||"):
            conteudo = parte.replace("||NEGRITO_INICIO||", "").replace("||NEGRITO_FIM||", "")
            run = paragrafo.add_run(conteudo)
            run.bold = True
        elif parte.startswith("||ITALICO_INICIO||") and parte.endswith("||ITALICO_FIM||"):
            conteudo = parte.replace("||ITALICO_INICIO||", "").replace("||ITALICO_FIM||", "")
            run = paragrafo.add_run(conteudo)
            run.italic = True
        else:
            run = paragrafo.add_run(parte)
        run.font.name = 'Book Antiqua'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Book Antiqua')

def substituir_runs_multiplo(paragrafo, substituicoes):
    texto_total = "".join(run.text for run in paragrafo.runs)

    # Substituições padrão
    for chave, valor in substituicoes.items():
        texto_total = texto_total.replace(chave, valor)

    # Negrito: <<NOME:...>> ou <<NEGRITO:...>>
    nomes = re.findall(r"<<NOME:(.*?)>>", texto_total)
    for nome in nomes:
        texto_total = texto_total.replace(f"<<NOME:{nome}>>", f"||NEGRITO_INICIO||{nome}||NEGRITO_FIM||")

    negritos = re.findall(r"<<NEGRITO:(.*?)>>", texto_total)
    for texto in negritos:
        texto_total = texto_total.replace(f"<<NEGRITO:{texto}>>", f"||NEGRITO_INICIO||{texto}||NEGRITO_FIM||")

    # Itálico fixo: aplicar à expressão "ad judicia et extra"
    texto_total = texto_total.replace("ad judicia et extra", "||ITALICO_INICIO||ad judicia et extra||ITALICO_FIM||")

    # Limpa runs e reinsere com estilo
    for run in paragrafo.runs:
        run.text = ""

    inserir_texto_com_estilo(paragrafo, texto_total)

def escolher_preposicao(referencia):
    """
    Decide 'na', 'no' ou 'em' com base na primeira palavra significativa da referência.
    A referência deve ser algo como: 'Rua ...', 'Avenida ...', 'Bairro ...', 'Brasília', etc.
    """
    if not referencia:
        return "em"

    # Normaliza: remove pontuação solta e números no começo
    ref = re.sub(r"^[\W_]+", "", referencia).strip()

    if not ref:
        return "em"

    # Se começar com uma abreviação, expande para o tipo de logradouro
    ref_low = ref.lower()
    abrevs = {
        "r.": "rua",
        "av.": "avenida",
        "rod.": "rodovia",
        "al.": "alameda",
        "tv.": "travessa",
        "pq.": "parque",
        "pç.": "praça",
        "estr.": "estrada",   # adicionado
    }
    # pega a primeira “palavra” visível
    primeira = ref_low.split()[0].rstrip(".,;:")

    if primeira in abrevs:
        tipo = abrevs[primeira]
    else:
        tipo = primeira

    # Listas de tipos com gênero gramatical comum em PT-BR
    femininas = {
        "rua", "avenida", "alameda", "travessa", "rodovia", "praça",
        "estrada", "via", "vila", "chácara", "fazenda", "comunidade"
    }
    masculinas = {
        "bairro", "condomínio", "loteamento", "setor", "residencial",
        "conjunto", "sítio", "parque", "resort"
    }

    if tipo in femininas:
        return "na"
    if tipo in masculinas:
        return "no"

    # Caso neutro/indefinido (cidades, estados, zonas não especificadas etc.)
    return "em"

# ========== Função principal de procuração ==========
def executar_procuracao_tool():
    try:
        quantidade = simpledialog.askinteger("Quantidade", "Quantas pessoas deseja incluir na procuração?")
        if not quantidade or quantidade <= 0:
            messagebox.showerror("Erro", "Quantidade inválida.")
            return

        pessoas = []
        for i in range(quantidade):
            messagebox.showinfo("Pessoa", f"Preenchendo dados da pessoa {i + 1}/{quantidade}")
            nome = simpledialog.askstring("Nome", "Nome completo:")
            genero = simpledialog.askstring("Gênero", "Gênero (M/F):")
            nacionalidade = simpledialog.askstring("Nacionalidade", "Nacionalidade:")
            estado_civil = simpledialog.askstring("Estado Civil", "Estado civil:")
            profissao = simpledialog.askstring("Profissão", "Profissão:")
            cpf = simpledialog.askstring("CPF", "CPF:")
            rg = simpledialog.askstring("RG", "RG:")
            rua = simpledialog.askstring("Rua", "Rua / Logradouro:")
            numero = simpledialog.askstring("Número", "Número:")
            bairro = simpledialog.askstring("Bairro", "Bairro:")
            cidade = simpledialog.askstring("Cidade", "Cidade:")
            cep = simpledialog.askstring("CEP", "CEP:")

            genero = (genero or "").strip().upper()
            is_feminino = genero == "F"

            palavra_portador = "portadora" if is_feminino else "portador"
            palavra_residente = "residente e domiciliada" if is_feminino else "residente e domiciliado"

            texto = f"<<NOME:{nome}>>"
            if nacionalidade:
                texto += f", {nacionalidade}"
            if estado_civil:
                texto += f", {estado_civil}"
            if profissao:
                texto += f", {profissao}"
            if cpf and rg:
                texto += f", {palavra_portador} do CPF nº {cpf} e RG nº {rg}"
            elif cpf:
                texto += f", {palavra_portador} do CPF nº {cpf}"
            elif rg:
                texto += f", {palavra_portador} do RG nº {rg}"

            # Monta endereço sem vírgulas sobrando
            partes_endereco = []
            if rua:
                partes_endereco.append(rua)
            if numero:
                partes_endereco.append(numero)
            if bairro:
                partes_endereco.append(bairro)
            if cidade:
                partes_endereco.append(cidade)

            endereco = ", ".join(partes_endereco)
            if cep:
                endereco = f"{endereco} - CEP: {cep}" if endereco else f"CEP: {cep}"

            # Remove espaços duplicados e aparas
            endereco = re.sub(r"\s{2,}", " ", endereco).strip()

            # Define a referência para a preposição (prioridade: rua > bairro > cidade)
            if rua:
                referencia_prep = rua
            elif bairro:
                referencia_prep = f"Bairro {bairro}"
            elif cidade:
                referencia_prep = cidade
            else:
                referencia_prep = ""  # cairá em 'em'

            preposicao = escolher_preposicao(referencia_prep)
            texto += f", {palavra_residente} {preposicao} {endereco}" if endereco else f", {palavra_residente}"
            pessoas.append(texto)

        # Proteção antes de usar pessoas[0]
        if not pessoas:
            messagebox.showerror("Erro", "Nenhuma pessoa foi informada.")
            return

        contato = simpledialog.askstring("WhatsApp", "Nome do contato no WhatsApp:")
        data_formatada = gerar_data_formatada()
        texto_identificacao = "; ".join(pessoas) + ","
        verbo_nomear = "nomeia" if quantidade == 1 else "nomeiam"
        verbo_constituir = "constitui" if quantidade == 1 else "constituem"

        substituicoes = {
            '<<TEXTO_IDENTIFICACAO>>': texto_identificacao,
            '<<VERBO_NOMEAR>>': verbo_nomear,
            '<<VERBO_CONSTITUIR>>': verbo_constituir,
            '<<DATA>>': data_formatada,
            '<<OUTORGANTE>>': "OUTORGANTE" if quantidade == 1 else "OUTORGANTES",
            '<<VERBO_PODER>>': "possa",
            '<<VERBO_CONFERE>>': "confere" if quantidade == 1 else "conferem",
            '<<PRONOME_POSSESSIVO>>': "seu" if quantidade == 1 else "seus",
            '<<PRONOME_OBJ>>': "o" if quantidade == 1 else "os",
            '<<AUTOR_RECLAMANTE>>': "autor ou reclamante" if quantidade == 1 else "autores ou reclamantes"
        }

        import sys
        if getattr(sys, 'frozen', False):
            base_path = sys._MEIPASS
        else:
            base_path = os.path.dirname(os.path.abspath(__file__))

        modelo_path = os.path.join(base_path, "PROCURACAO MODELO1.docx")
        pasta_destino = os.path.join(base_path, "documentos_gerados")
        os.makedirs(pasta_destino, exist_ok=True)

        nome_base = pessoas[0].split(",")[0].replace("<<NOME:", "").replace(">>", "").strip().replace(" ", "_")
        docx_path = os.path.join(pasta_destino, f"PROCURACAO_{nome_base}.docx")
        pdf_path = os.path.join(pasta_destino, f"PROCURACAO_{nome_base}.pdf")

        documento = Document(modelo_path)
        for p in documento.paragraphs:
            substituir_runs_multiplo(p, substituicoes)
        for tabela in documento.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for p in celula.paragraphs:
                        substituir_runs_multiplo(p, substituicoes)

        documento.save(docx_path)
        convert(docx_path, pdf_path)

        if messagebox.askyesno("WhatsApp", "Deseja enviar o PDF via WhatsApp Web?"):
            enviar_pdf_pyautogui(contato, pdf_path)

        messagebox.showinfo("Concluído", f"Documentos gerados em:\n{pasta_destino}")

    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro:\n{e}")

def enviar_pdf_pyautogui(nome_contato, caminho_pdf):
    webbrowser.open('https://web.whatsapp.com')
    time.sleep(15)
    pyautogui.click(313, 193)
    time.sleep(1)
    pyperclip.copy(nome_contato)
    pyautogui.hotkey('ctrl', 'v')
    time.sleep(3)
    pyautogui.press('enter')
    time.sleep(2)
    pyautogui.click(710, 966)
    time.sleep(1)
    pyautogui.click(690, 531)
    time.sleep(1)
    pyperclip.copy(caminho_pdf)
    pyautogui.hotkey('ctrl', 'v')
    time.sleep(1)
    pyautogui.press('enter')
    time.sleep(2)
    pyautogui.press('enter')
    time.sleep(3)
    pyautogui.hotkey('ctrl', 'w')

# ========== Interface Gráfica ==========
def main():
    root = tk.Tk()
    root.title("Ferramentas de Escritório")
    root.geometry("420x250")
    root.resizable(False, False)
    try:
        root.iconbitmap("OIP.ico")
    except:
        pass
    tk.Label(root, text="Selecione a ferramenta que deseja utilizar:", font=("Segoe UI", 12)).pack(pady=20)
    tk.Button(root, text="📁 Unir e Comprimir PDFs", command=executar_pdf_tool, width=30, height=2).pack(pady=10)
    tk.Button(root, text="📄 Gerar e Enviar Procuração", command=executar_procuracao_tool, width=30, height=2).pack(pady=10)
    root.mainloop()

if __name__ == "__main__":
    main()
