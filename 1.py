import os
import re
import sys
import pandas as pd
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import time

# Funções auxiliares
def valor_str(valor):
    return str(valor).strip() if pd.notna(valor) else ""

def formatar_nome_proprio(texto):
    return " ".join(w.capitalize() for w in texto.split())

def extrair_artigos_de_lei(texto):
    padrao = r"(art\.?\s*\d+º?(?:[-–]?\s*\d+)?\s*(?:do\s+(?:CPC|Código Civil|CC|CF|CLT|CP|CPP|CDC|ECA|Lei nº?\s*\d+/\d+)))"
    artigos = re.findall(padrao, texto, flags=re.IGNORECASE)
    artigos_limpos = sorted(set(a.strip().capitalize() for a in artigos))
    return artigos_limpos

# Inicialização
temas_excluidos = []
fundamentos_exclusao = ""
pedidos_exclusao = ""

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

arquivo_excel = r"C:\\Users\\lealf\\OneDrive\\Área de trabalho - atual\\dados_peticao.xlsx"
arquivo_modelo = r"C:\\Users\\lealf\\OneDrive\\Área de trabalho - atual\\MODELO.docx"
arquivo_saida = r"C:\\Users\\lealf\\OneDrive\\Área de trabalho - atual\\peticao_gerada1.docx"

partes_df = pd.read_excel(arquivo_excel, sheet_name="Partes")
processo_df = pd.read_excel(arquivo_excel, sheet_name="Processo")
advogado_df = pd.read_excel(arquivo_excel, sheet_name="Advogado")
try:
    secoes_opcionais_df = pd.read_excel(arquivo_excel, sheet_name="SecoesOpcionais")
    temas_excluidos = [valor_str(t).lower() for t in secoes_opcionais_df["Título"] if pd.notna(t)]
except Exception:
    secoes_opcionais_df = pd.DataFrame()

# Dados do processo
processo = processo_df.iloc[0]
comarca_completa = valor_str(processo["Comarca"])
comarca = re.sub(r"\s*[-–]?\s*paraíba\s*$", "", comarca_completa, flags=re.IGNORECASE).strip()
vara = valor_str(processo["Vara"]).upper()
tipo_causa = valor_str(processo["Tipo de Causa"])
fatos = valor_str(processo["Fatos"])
valor_causa = processo["Valor da Causa"]

adv = advogado_df.iloc[0]
nome_adv = valor_str(adv['Nome Completo']).upper()

requerentes = partes_df[partes_df['Tipo'].str.lower().str.strip() == 'requerente']
requeridos = partes_df[partes_df['Tipo'].str.lower().str.strip() == 'requerido']

def gerar_qualificacao(pessoas):
    textos = []
    for _, pessoa in pessoas.iterrows():
        nome = valor_str(pessoa['Nome Completo']).upper()
        elementos = []
        for campo, rotulo in [
            ('Nacionalidade', ''),
            ('Estado Civil', ''),
            ('Profissão', ''),
            ('CPF', 'portador(a) do CPF nº {}'),
            ('RG', 'e RG nº {}'),
            ('Endereço', 'residente e domiciliado(a) na(o) {}'),
            ('Bairro', 'Bairro {}'),
            ('Cidade', ''),
            ('CEP', 'CEP {}')
        ]:
            # Campo Cidade tratado em conjunto com UF
            if campo == 'Cidade' and pessoa.get('UF'):
                cidade = valor_str(pessoa['Cidade'])
                uf = valor_str(pessoa['UF']).upper()
                if cidade:
                    elementos.append(f"{cidade}/{uf}")
                continue

            valor = valor_str(pessoa[campo]) if campo in pessoa else ''
            if not valor:
                continue

            # Tratamento especial para RG
            if campo == 'RG':
                partes = valor.rsplit(' ', 1)
                if len(partes) == 2:
                    numero, orgao = partes[0], partes[1].upper()
                    valor_corrigido = f"{numero} {orgao}"
                else:
                    valor_corrigido = valor.upper()
                elementos.append(f"e RG nº {valor_corrigido}")

            # Campos que devem permanecer em minúsculo
            elif campo in ['Nacionalidade', 'Estado Civil', 'Profissão']:
                elementos.append(valor.lower())

            else:
                elementos.append(
                    rotulo.format(formatar_nome_proprio(valor)) if '{}' in rotulo else formatar_nome_proprio(valor))

        textos.append((nome, ", ".join(elementos)))
    return textos

qualificacoes_requerentes = gerar_qualificacao(requerentes)
qualificacoes_requeridos = gerar_qualificacao(requeridos)

# Documento
documento = Document(arquivo_modelo)

def adicionar_paragrafo(texto, negrito=False, maiusculo=False, tamanho=12, centralizado=False, titulo=False):
    texto = re.sub(r'\s+', ' ', texto.replace('\n', ' ')).strip()
    if maiusculo:
        texto = texto.upper()
    par = documento.add_paragraph()
    run = par.add_run(texto)
    run.font.name = 'Book Antiqua'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Book Antiqua')
    run.font.size = Pt(tamanho)
    run.bold = negrito
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER if centralizado else WD_ALIGN_PARAGRAPH.JUSTIFY
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.first_line_indent = Cm(0 if titulo else 1.25)
    return par

# Título e qualificação
adicionar_paragrafo(f"EXCELENTÍSSIMA SENHORA DOUTORA JUIZA DE DIREITO DA {vara} DA COMARCA DE {comarca} - PARAÍBA", negrito=True, maiusculo=True, tamanho=16, centralizado=True, titulo=True)
for _ in range(3):
    adicionar_paragrafo("")

tipo_acao_formatado = tipo_causa.upper() if tipo_causa.upper().startswith("AÇÃO DE") else f"AÇÃO DE {tipo_causa.upper()}"

par = documento.add_paragraph()
par.paragraph_format.line_spacing = 1.5
par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for idx, (nome, qual) in enumerate(qualificacoes_requerentes):
    run_nome = par.add_run(nome)
    run_nome.bold = True
    run_nome.font.name = 'Book Antiqua'
    run_nome._element.rPr.rFonts.set(qn('w:eastAsia'), 'Book Antiqua')
    run_nome.font.size = Pt(12)

    run_qual = par.add_run(", " + qual)
    run_qual.font.name = 'Book Antiqua'
    run_qual._element.rPr.rFonts.set(qn('w:eastAsia'), 'Book Antiqua')
    run_qual.font.size = Pt(12)

    if idx < len(qualificacoes_requerentes) - 1:
        run_sep = par.add_run("; ")
        run_sep.font.name = 'Book Antiqua'
        run_sep._element.rPr.rFonts.set(qn('w:eastAsia'), 'Book Antiqua')
        run_sep.font.size = Pt(12)
    else:
        # Texto antes do nome do advogado
        run_intro = par.add_run("; neste ato representado por ")
        run_intro.font.name = 'Book Antiqua'
        run_intro._element.rPr.rFonts.set(qn('w:eastAsia'), 'Book Antiqua')
        run_intro.font.size = Pt(12)

        # Nome do advogado em negrito
        run_adv = par.add_run(nome_adv)
        run_adv.bold = True
        run_adv.font.name = 'Book Antiqua'
        run_adv._element.rPr.rFonts.set(qn('w:eastAsia'), 'Book Antiqua')
        run_adv.font.size = Pt(12)

        # Texto intermediário
        run_oab = par.add_run(f", inscrito na OAB/{valor_str(adv['UF']).upper()} sob o nº {valor_str(adv['OAB'])}, vem, respeitosamente, à presença de Vossa Excelência propor a presente ")
        run_oab.font.name = 'Book Antiqua'
        run_oab._element.rPr.rFonts.set(qn('w:eastAsia'), 'Book Antiqua')
        run_oab.font.size = Pt(12)

        # Nome da ação em negrito
        run_acao = par.add_run(tipo_acao_formatado)
        run_acao.bold = True
        run_acao.font.name = 'Book Antiqua'
        run_acao._element.rPr.rFonts.set(qn('w:eastAsia'), 'Book Antiqua')
        run_acao.font.size = Pt(12)

# Requerido (na mesma linha)
if qualificacoes_requeridos:
    req_texto = ", ".join([f"{nome}, {qual}" for nome, qual in qualificacoes_requeridos])
    run_req = par.add_run(f", em face de {req_texto}, pelos fatos e fundamentos que passa a expor:")
    run_req.font.name = 'Book Antiqua'
    run_req._element.rPr.rFonts.set(qn('w:eastAsia'), 'Book Antiqua')
    run_req.font.size = Pt(12)
else:
    run_fatos = par.add_run(", pelos fatos e fundamentos que passa a expor:")
    run_fatos.font.name = 'Book Antiqua'
    run_fatos._element.rPr.rFonts.set(qn('w:eastAsia'), 'Book Antiqua')
    run_fatos.font.size = Pt(12)

# Reescrita dos fatos
fatos_interpretados = client.chat.completions.create(
    model="gpt-4o",
    messages=[{"role": "user", "content": f"Você é um advogado especialista. Reescreva os fatos da ação de {tipo_causa}:\n{fatos}"}],
    temperature=0.2,
    max_tokens=1500
).choices[0].message.content.strip()

adicionar_paragrafo("")
adicionar_paragrafo("1 - DOS FATOS", negrito=True, maiusculo=True, tamanho=14, titulo=True)
adicionar_paragrafo("")
for par in fatos_interpretados.splitlines():
    if par.strip():
        adicionar_paragrafo(par.strip(), tamanho=12)

# Gratuidade da Justiça
print("Deseja incluir a seção de gratuidade da justiça? (sim/não)")
incluir_gratuidade = input().strip().lower() == 'sim'
if incluir_gratuidade:
    texto_bruto = client.chat.completions.create(
        model="gpt-4o",
        messages=[{
            "role": "user",
            "content": (
                "Você é um advogado especialista. Escreva parágrafos formais, claros e fundamentados sobre o pedido de gratuidade de justiça, "
                "com base na Constituição Federal e Código de Processo Civil, para ser usado em petição judicial. "
                "Não inclua introduções genéricas nem conclua com frases como 'pode ser adaptado'."
            )
        }],
        temperature=0.3,
        max_tokens=600
    ).choices[0].message.content.strip()

    def limpar_conteudo_gerado(texto):
        linhas = texto.strip().splitlines()
        corpo = [linha.strip() for linha in linhas if linha.strip()]
        introducoes = [
            "certamente", "a seguir", "segue", "veja", "com base", "conforme solicitado",
            "apresento", "texto pode", "a título", "observe", "exemplo", "modelo", "abaixo"
        ]
        conclusoes = [
            "este texto", "pode ser adaptado", "conforme o caso", "observe", "exemplo genérico",
            "segue exemplo", "caso concreto", "pede deferimento", "termos em que"
        ]
        corpo_limpo = []
        for linha in corpo:
            linha_min = linha.lower()
            if any(linha_min.startswith(i) for i in introducoes):
                continue
            if any(c in linha_min for c in conclusoes):
                continue
            corpo_limpo.append(linha)
        return "\n".join(corpo_limpo).strip()

    texto_gratuidade = limpar_conteudo_gerado(texto_bruto)

    adicionar_paragrafo("")
    adicionar_paragrafo("2 - DA GRATUIDADE DA JUSTIÇA", negrito=True, maiusculo=True, tamanho=14, titulo=True)
    adicionar_paragrafo("")
    for par in texto_gratuidade.splitlines():
        if par.strip():
            adicionar_paragrafo(par.strip(), tamanho=12)

# Geração dos fundamentos
fundamentos_gerados = client.chat.completions.create(
    model="gpt-4o",
    messages=[{
        "role": "user",
        "content": (
            f"Você é um advogado especialista. Com base nos fatos a seguir de uma ação de {tipo_causa}, "
            f"gere uma seção de fundamentos jurídicos com base apenas em legislação brasileira, como a Constituição Federal, "
            f"o Código de Processo Civil e outras leis pertinentes.{fundamentos_exclusao} Não utilize jurisprudência ou doutrina. "
            f"Utilize linguagem formal, clara e precisa.\n\n"
            f"Fatos: {fatos}\n\n"
            f"Apenas gere o conteúdo dos fundamentos, não repita o título. Apresente apenas os parágrafos. Evite incluir no documento algo que não seja os argumentos sobre a gratuidade."
        )
    }],
    temperature=0.2,
    max_tokens=1800
).choices[0].message.content.strip()



adicionar_paragrafo("")
adicionar_paragrafo("3 - DOS FUNDAMENTOS JURÍDICOS", negrito=True, maiusculo=True, tamanho=14, titulo=True)
adicionar_paragrafo("")
for par in fundamentos_gerados.splitlines():
    if par.strip():
        adicionar_paragrafo(par.strip(), tamanho=12)

# Seções opcionais
print("Deseja incluir as seções opcionais da aba 'SecoesOpcionais'? (sim/não)")
incluir_secoes_opcionais = input().strip().lower() == 'sim'

indice_secao = 4
def titulo_secao(titulo):
    global indice_secao
    adicionar_paragrafo("")
    adicionar_paragrafo(f"{indice_secao} - {titulo.upper()}", negrito=True, maiusculo=True, tamanho=14, titulo=True)
    adicionar_paragrafo("")
    indice_secao += 1

if incluir_secoes_opcionais and not secoes_opcionais_df.empty:
    for _, linha in secoes_opcionais_df.iterrows():
        titulo_secao_bruto = valor_str(linha.get("Título"))
        observacao = valor_str(linha.get("Observação"))
        if not titulo_secao_bruto:
            continue
        prompt = (
            f"Você é um advogado especialista. Redija a seção intitulada '{titulo_secao_bruto}', "
            f"com base na legislação brasileira e nas observações fornecidas: \"{observacao}\". "
            f"Considere que a ação trata de {tipo_causa}. Use linguagem formal e estruturada. Não repita o título."
        )
        try:
            conteudo_secao = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
                max_tokens=1000
            ).choices[0].message.content.strip()

            titulo_secao(titulo_secao_bruto)
            for par in conteudo_secao.splitlines():
                if par.strip() and not par.strip().upper().startswith(titulo_secao_bruto.upper()):
                    adicionar_paragrafo(par.strip(), tamanho=12)

        except Exception as e:
            print(f"Erro ao gerar seção '{titulo_secao_bruto}': {e}")

# Geração dos pedidos
pedidos_gerados = client.chat.completions.create(
    model="gpt-4o",
    messages=[{
        "role": "user",
        "content": (
            f"Você é um advogado especialista. Com base nos seguintes fatos de uma ação de {tipo_causa}, "
            f"elabore a seção de PEDIDOS da petição inicial. Liste cada pedido em tópico separado, utilizando linguagem formal e clara. "
            f"Inclua o fundamento legal aplicável (artigo de lei, CPC, CF, CLT etc.) somente quando necessário. "
            f"Evite repetir fundamentos legais desnecessários.\n\n"
            f"Fatos: {fatos}\n\n"
            f"Responda apenas com os pedidos, em tópicos, sem introdução, sem repetir os fatos e sem rodapé.{pedidos_exclusao}"
        )
    }],
    temperature=0.2,
    max_tokens=1500
).choices[0].message.content.strip()

titulo_secao("Dos Pedidos")
adicionar_paragrafo("Diante do exposto, requerem a Vossa Excelência:", tamanho=12, titulo=True)
adicionar_paragrafo("")

pedidos_limpos = "\n".join([
    re.sub(r"^\(?[a-zA-Z0-9]+\)?[.\)]?\s*", "", linha.strip())
    for linha in pedidos_gerados.split("\n") if linha.strip()
])
for i, linha in enumerate(pedidos_limpos.splitlines()):
    letra = chr(97 + i)
    adicionar_paragrafo(f"{letra}) {linha}", tamanho=12)

# Encerramento
adicionar_paragrafo("")
valor_formatado = f"{valor_causa:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
adicionar_paragrafo(f"Dá-se à presente causa o valor de R$ {valor_formatado}", tamanho=12, titulo=True)
adicionar_paragrafo("")
adicionar_paragrafo("Nesses termos,", tamanho=12, titulo=True)
adicionar_paragrafo("Pede deferimento.", tamanho=12, titulo=True)
adicionar_paragrafo("")

# Assinatura
adicionar_paragrafo(formatar_nome_proprio(nome_adv), tamanho=12, negrito=True, titulo=True)
adicionar_paragrafo(f"OAB/{valor_str(adv['UF']).upper()} {valor_str(adv['OAB'])}", tamanho=12, negrito=True, titulo=True)

# Salvar
for tentativa in range(3):
    try:
        documento.save(arquivo_saida)
        print(f"✅ Petição gerada com sucesso em {arquivo_saida}")
        break
    except PermissionError:
        print(f"❌ Arquivo está aberto. Feche e tentarei novamente ({tentativa + 1}/3)...")
        time.sleep(5)
else:
    print("❌ Não consegui salvar. Verifique se o arquivo está fechado.")
