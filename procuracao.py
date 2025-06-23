from docx import Document
from datetime import datetime
import win32com.client
import os
import time
import webbrowser
import pyautogui
import pyperclip

def gerar_data_formatada():
    meses = {
        1: 'janeiro', 2: 'fevereiro', 3: 'março',
        4: 'abril', 5: 'maio', 6: 'junho',
        7: 'julho', 8: 'agosto', 9: 'setembro',
        10: 'outubro', 11: 'novembro', 12: 'dezembro'
    }
    hoje = datetime.now()
    return f"{hoje.day} de {meses[hoje.month]} de {hoje.year}"

def substituir_runs(paragrafo, substituicoes, nome_chave='<<NOME_COMPLETO>>', nome_valor=None):
    runs = paragrafo.runs
    i = 0
    while i < len(runs):
        run = runs[i]
        texto = run.text

        # Substituição especial do nome (com negrito mantendo fonte original)
        if nome_chave in texto and nome_valor is not None:
            partes = texto.split(nome_chave)

            # Substituir o texto atual pelo antes do nome
            run.text = partes[0]

            # Criar run para o nome em negrito mantendo a formatação do run original
            negrito_run = paragrafo.add_run(nome_valor)
            negrito_run.bold = True
            negrito_run.font.name = run.font.name
            negrito_run.font.size = run.font.size
            negrito_run.font.color.rgb = run.font.color.rgb
            negrito_run.font.italic = run.font.italic
            negrito_run.font.underline = run.font.underline

            # Criar run para o texto depois do nome
            depois_run = paragrafo.add_run(partes[1])

            # Inserir os novos runs na posição correta
            paragrafo._p.insert(paragrafo._p.index(run._r) + 1, negrito_run._r)
            paragrafo._p.insert(paragrafo._p.index(negrito_run._r) + 1, depois_run._r)

            i += 3
            continue

        # Substituir outras tags mantendo formatação original do run
        for chave, valor in substituicoes.items():
            if chave in texto:
                run.text = texto.replace(chave, valor)
                texto = run.text  # atualiza texto

        i += 1

def substituir_no_documento(doc, substituicoes, nome_chave='<<NOME_COMPLETO>>', nome_valor=None):
    for p in doc.paragraphs:
        substituir_runs(p, substituicoes, nome_chave, nome_valor)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for p in celula.paragraphs:
                    substituir_runs(p, substituicoes, nome_chave, nome_valor)

def converter_docx_para_pdf(caminho_docx, caminho_pdf):
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(caminho_docx)
    doc.SaveAs(caminho_pdf, FileFormat=17)  # 17 = PDF
    doc.Close()
    word.Quit()

def enviar_pdf_pyautogui(nome_contato, caminho_pdf):
    print("🟢 Abrindo WhatsApp Web...")
    webbrowser.open('https://web.whatsapp.com')
    time.sleep(15)  # Tempo para login no WhatsApp Web

    print("⌛ Procurando contato...")
    pyautogui.click(313, 193)  # Ajuste para sua tela (campo de busca)
    time.sleep(1)
    pyperclip.copy(nome_contato)
    pyautogui.hotkey('ctrl', 'v')
    time.sleep(3)

    pyautogui.click(248, 427)  # Clique no contato (ajuste para sua tela)
    time.sleep(2)

    print("⌛ Clicando no clipe...")
    pyautogui.click(710, 966)  # Ícone do clipe (ajuste para sua tela)
    time.sleep(1)

    print("⌛ Clicando em Documento...")
    pyautogui.click(690, 531)  # Documento (ajuste para sua tela)
    time.sleep(1)

    print("⌛ Inserindo caminho...")
    pyperclip.copy(caminho_pdf)
    pyautogui.hotkey('ctrl', 'v')
    time.sleep(1)

    pyautogui.press('enter')
    time.sleep(2)

    print("⌛ Enviando arquivo...")
    pyautogui.press('enter')
    time.sleep(3)

    print("❌ Fechando a aba do WhatsApp Web...")
    pyautogui.hotkey('ctrl', 'w')

    print("✅ Arquivo enviado com sucesso e aba fechada!")

def main():
    nome = input("Nome completo: ").strip()
    nacionalidade = input("Nacionalidade: ").strip()
    estado_civil = input("Estado civil: ").strip()
    profissao = input("Profissão: ").strip()
    cpf = input("CPF (000.000.000-00): ").strip()
    rg = input("RG (Ex.: 3.652.546 SSP/PB): ").strip()
    endereco = input("Rua / Logradouro: ").strip()
    numero = input("Número da casa: ").strip()
    bairro = input("Bairro: ").strip()
    cidade = input("Cidade (Ex.: Princesa Isabel/PB): ").strip()
    cep = input("CEP (ou deixe em branco): ").strip()
    contato = input("Nome do contato no WhatsApp (como está salvo): ").strip()

    data_formatada = gerar_data_formatada()

    profissao_texto = f", {profissao}" if profissao else ""
    numero_texto = f", nº {numero}" if numero else ""
    cep_texto = f" - CEP: {cep}" if cep else ""

    substituicoes = {
        '<<NACIONALIDADE>>': nacionalidade,
        '<<ESTADO_CIVIL>>': estado_civil,
        '<<PROFISSAO>>': profissao_texto,
        '<<CPF>>': cpf,
        '<<RG>>': rg,
        '<<ENDERECO>>': endereco,
        '<<NUMERO>>': numero_texto,
        '<<BAIRRO>>': bairro,
        '<<CIDADE>>': cidade,
        '<<CEP>>': cep_texto,
        '<<DATA>>': data_formatada,
    }

    base_path = r"C:\Users\lealf\OneDrive\Área de trabalho - atual\ARQUIVOS ESCRITÓRIO DE EMANUEL\PROCURACOES GERADAS"
    nome_arquivo = f"PROCURACAO_{nome.replace(' ', '_')}"
    caminho_modelo = os.path.join(base_path, "PROCURACAO MODELO1.docx")
    caminho_docx = os.path.join(base_path, f"{nome_arquivo}.docx")
    caminho_pdf = os.path.join(base_path, f"{nome_arquivo}.pdf")

    documento = Document(caminho_modelo)
    substituir_no_documento(documento, substituicoes, nome_chave='<<NOME_COMPLETO>>', nome_valor=nome)
    documento.save(caminho_docx)
    print("📝 Documento preenchido com sucesso!")

    converter_docx_para_pdf(caminho_docx, caminho_pdf)
    print("📄 Documento convertido em PDF com sucesso!")

    enviar_pdf_pyautogui(contato, caminho_pdf)

if __name__ == "__main__":
    main()
