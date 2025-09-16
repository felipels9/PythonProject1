import os
import re
import time
import json
import random
import pathlib
from typing import List, Tuple, Optional

import pandas as pd
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# =========================
# CONFIGURAÇÃO DO USUÁRIO
# =========================
ARQUIVO_EXCEL   = r"C:\Users\lealf\OneDrive\Área de trabalho - atual\dados_peticao.xlsx"
ARQUIVO_MODELO  = r"C:\Users\lealf\OneDrive\Área de trabalho - atual\MODELO.docx"
ARQUIVO_SAIDA   = r"C:\Users\lealf\OneDrive\Área de trabalho - atual\peticao_gerada1.docx"

INCLUIR_GRATUIDADE              = True
INCLUIR_SECOES_OPCIONAIS        = True
INCLUIR_SUMARIO_DISPOSITIVOS    = True   # lista “art. ... do CPC/CF/...” encontrados ao final
TEMPO_ESPERA_SALVAR_SEG         = 5
TENTATIVAS_SALVAR               = 3

# =========================
# TEXTOS / FORMATADORES
# =========================
def valor_str(valor) -> str:
    return str(valor).strip() if pd.notna(valor) and str(valor).strip() != "nan" else ""

def limpar_espacos(texto: str) -> str:
    return re.sub(r"\s+", " ", str(texto).replace("\n", " ")).strip()

def formatar_nome_proprio(texto: str) -> str:
    return " ".join(w.capitalize() for w in str(texto).split())

def extrair_artigos_de_lei(texto: str) -> List[str]:
    padrao = r"(art\.?\s*\d+º?(?:[-–]?\s*\d+)?\s*(?:do\s+(?:CPC|Código Civil|CC|CF|CLT|CP|CPP|CDC|ECA|Lei\s*n[ºo]?\s*\d+/\d+)))"
    artigos = re.findall(padrao, texto, flags=re.IGNORECASE)
    artigos = [limpar_espacos(a) for a in artigos]
    return sorted(set(artigos), key=lambda x: x.lower())

def moeda_brasil(valor: Optional[float]) -> str:
    try:
        v = float(valor) if valor not in (None, "",) else 0.0
    except Exception:
        v = 0.0
    s = f"{v:,.2f}"
    return s.replace(",", "X").replace(".", ",").replace("X", ".")

# =========================
# OPENAI (com fallback)
# =========================
def gpt_or_fallback(prompt: str, fallback: str, max_tokens: int = 1200, temperature: float = 0.2) -> str:
    """
    Tenta usar OpenAI; se indisponível (sem chave/erro), devolve o fallback.
    """
    try:
        from openai import OpenAI
        load_dotenv()
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            return fallback
        client = OpenAI(api_key=api_key)
        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=temperature,
            max_tokens=max_tokens
        )
        texto = resp.choices[0].message.content.strip()
        return texto if texto else fallback
    except Exception:
        return fallback

# =========================
# HISTÓRICO + SIMILARIDADE
# =========================
HIST_PATH = pathlib.Path(ARQUIVO_EXCEL).with_name("historico_textos.jsonl")

def _carregar_historico():
    itens = []
    if HIST_PATH.exists():
        with open(HIST_PATH, "r", encoding="utf-8") as f:
            for line in f:
                try:
                    itens.append(json.loads(line))
                except Exception:
                    continue
    return itens

def _salvar_historico(secao_key: str, texto: str):
    HIST_PATH.parent.mkdir(parents=True, exist_ok=True)
    registro = {"secao": secao_key, "texto": texto, "ts": time.time()}
    with open(HIST_PATH, "a", encoding="utf-8") as f:
        f.write(json.dumps(registro, ensure_ascii=False) + "\n")

def _ngram_set(texto: str, n: int = 3):
    tokens = re.findall(r"\w+|\S", texto.lower())
    return {" ".join(tokens[i:i+n]) for i in range(max(0, len(tokens)-n+1))}

def similaridade_jaccard(a: str, b: str, n: int = 3) -> float:
    A = _ngram_set(a, n=n)
    B = _ngram_set(b, n=n)
    if not A or not B:
        return 0.0
    inter = len(A & B)
    union = len(A | B)
    return inter / union if union else 0.0

def gerar_texto_unico(
    secao_key: str,
    construir_prompt_fn,
    fatos_contexto: str,
    tentativas: int = 6,
    alvo_max_similaridade: float = 0.33,   # quanto MENOR, mais diferente; 0.25~0.45 recomendado
    ngram_n: int = 3,
    temperaturas = (0.4, 0.6, 0.8, 0.9, 1.0, 1.1),
    variantes_instrucao=None,
    fallback_padrao: str = ""
) -> str:
    """
    Gera texto para uma 'secao_key' tentando ficar distinto do histórico.
    - construir_prompt_fn: (fatos_contexto, instrucao_variante) -> prompt
    - variantes_instrucao: lista de estilos/ênfases para rotacionar
    """
    historico = _carregar_historico()
    corpus_secao = [h["texto"] for h in historico if h.get("secao") == secao_key]
    variantes = list(variantes_instrucao or [""])
    random.shuffle(variantes)

    melhor_texto = ""
    menor_sim = 1.0

    for i in range(tentativas):
        instr = variantes[i % len(variantes)]
        temp = temperaturas[i % len(temperaturas)]
        prompt = construir_prompt_fn(fatos_contexto, instr)

        texto = gpt_or_fallback(prompt, fallback="", max_tokens=1100, temperature=temp).strip()
        if not texto:
            continue

        sim_maior = 0.0
        for antigo in corpus_secao:
            sim = similaridade_jaccard(texto, antigo, n=ngram_n)
            if sim > sim_maior:
                sim_maior = sim

        if sim_maior < menor_sim:
            menor_sim = sim_maior
            melhor_texto = texto

        if sim_maior <= alvo_max_similaridade:
            _salvar_historico(secao_key, texto)
            return texto

    if melhor_texto:
        _salvar_historico(secao_key, melhor_texto)
    return melhor_texto or fallback_padrao or " "

# =========================
# VARIANTES DE ESTILO
# =========================
# — Gratuidade
VARIANTES_GRATUIDADE = [
    ("Redija a seção de justiça gratuita com base no art. 5º, LXXIV, da CF e arts. 98-102 do CPC, "
     "variando conectivos e sinonímias; evite clichês e fórmulas previsíveis; sem título e sem rodapé."),
    ("Elabore o pedido de gratuidade enfatizando acesso à justiça (art. 5º, XXXV, CF) e presunção da declaração "
     "de hipossuficiência (art. 99, §3º, CPC), trocando verbos e ordem lógica dos argumentos."),
    ("Escreva destacando que despesas não podem obstar a tutela jurisdicional (CF, art. 5º, LXXIV), "
     "com registro do regime do CPC e do ônus de impugnação por prova em contrário; linguagem técnica e concisa."),
    ("Produza com foco na proporcionalidade e efetividade da tutela, evitando construções repetidas, "
     "e mantendo tom jurídico impessoal e objetivo, sem frases de efeito."),
]

def montar_prompt_gratuidade(fatos_resumidos: str, instrucao: str) -> str:
    return (
        f"Você é advogado especialista. {instrucao} "
        f"Se útil, considere este contexto: {fatos_resumidos} "
        f"Entregue apenas parágrafos da seção, sem títulos, tópicos ou rodapés."
    )

# — Fundamentos
VARIANTES_FUNDAMENTOS = [
    ("Redija fundamentos jurídicos somente com legislação (CF, CPC, leis correlatas), "
     "alternando ordem dos dispositivos e conectivos, e evitando padronização; sem doutrina, sem jurisprudência."),
    ("Escreva fundamentos destacando princípios constitucionais (isonomia, devido processo, inafastabilidade) "
     "e regras instrumentais do CPC, variando o vocabulário técnico e a arquitetura frasal."),
    ("Produza fundamentos enfatizando deveres do juiz (arts. 4º, 6º, 139 CPC), contraditório substancial e "
     "cooperação processual; evite sequências cliché e repetições de trechos."),
    ("Formule fundamentos com foco em tutela adequada e efetiva, procedimento legal e segurança jurídica, "
     "com linguagem enxuta e sem redundância, variando a estrutura dos períodos."),
]

def montar_prompt_fundamentos(fatos_ctx: str, instrucao: str, tipo_causa: str) -> str:
    return (
        f"Você é advogado especialista. {instrucao} "
        f"Contexto (ação de {tipo_causa}): {fatos_ctx} "
        f"Responda apenas com parágrafos (sem título, sem tópicos)."
    )

# — Pedidos
VARIANTES_PEDIDOS = [
    ("Liste os pedidos em linhas separadas, linguagem objetiva, indicando fundamento legal apenas quando indispensável; "
     "varie verbos de requerimento e a ordem dos itens; não use introdução ou rodapé."),
    ("Elabore pedidos sucintos e completos, com alternância de conectivos e estruturas, sem frases padronizadas; "
     "evite repetição de fórmulas e mantenha técnica; apenas os itens, um por linha."),
    ("Produza a seção de pedidos com foco em clareza, precisão e suficiência, trocando a redação costumeira e "
     "evitando perífrases comuns; cada pedido em uma linha."),
    ("Formule pedidos com foco na utilidade prática (citação, procedência, condenações, custas/honorários), "
     "variando léxico e sintaxe; sem preâmbulo e sem fecho."),
]

def montar_prompt_pedidos(fatos_ctx: str, instrucao: str, tipo_causa: str) -> str:
    return (
        f"Você é advogado especialista. {instrucao} "
        f"Baseie-se nestes fatos (ação de {tipo_causa}): {fatos_ctx} "
        f"Entregue somente os itens, um por linha."
    )

# =========================
# EXCEL → DATAFRAMES
# =========================
def ler_planilhas(caminho: str):
    try:
        partes_df   = pd.read_excel(caminho, sheet_name="Partes").fillna("")
        processo_df = pd.read_excel(caminho, sheet_name="Processo").fillna("")
        advogado_df = pd.read_excel(caminho, sheet_name="Advogado").fillna("")
    except Exception as e:
        raise RuntimeError(f"Erro ao ler Excel '{caminho}': {e}")

    try:
        secoes_opcionais_df = pd.read_excel(caminho, sheet_name="SecoesOpcionais").fillna("")
    except Exception:
        secoes_opcionais_df = pd.DataFrame(columns=["Título", "Observação"])

    return partes_df, processo_df, advogado_df, secoes_opcionais_df

# =========================
# QUALIFICAÇÃO (com gênero)
# =========================
def gerar_qualificacao(pessoas: pd.DataFrame) -> List[Tuple[str, str]]:
    """
    Espera colunas (quando houver): Nome Completo, Genero (M/F), Nacionalidade, Estado Civil, Profissão,
    CPF, RG (p.ex. '123456 SSP-PB' / '123456 SSP PB'), Endereço, Bairro, Cidade, UF, CEP.
    """
    textos = []
    for _, p in pessoas.iterrows():
        nome = valor_str(p.get("Nome Completo")).upper()

        genero = valor_str(p.get("Genero")).upper()
        fem = (genero == "F")
        portador = "portadora" if fem else "portador"
        domiciliado = "domiciliada" if fem else "domiciliado"

        elementos = []

        for campo in ("Nacionalidade", "Estado Civil", "Profissão"):
            v = valor_str(p.get(campo))
            if v:
                elementos.append(v.lower())

        cpf = valor_str(p.get("CPF"))
        if cpf:
            elementos.append(f"{portador} do CPF nº {cpf}")

        rg = valor_str(p.get("RG"))
        if rg:
            rg_limpo = limpar_espacos(rg)
            partes = rg_limpo.rsplit(" ", 1)
            if len(partes) == 2 and len(partes[1]) <= 10:
                numero, orgao = partes[0], partes[1].upper()
                elementos.append(f"e RG nº {numero} {orgao}")
            else:
                elementos.append(f"e RG nº {rg_limpo.upper()}")

        endereco = valor_str(p.get("Endereço"))
        bairro   = valor_str(p.get("Bairro"))
        cidade   = valor_str(p.get("Cidade"))
        uf       = valor_str(p.get("UF")).upper()
        cep      = valor_str(p.get("CEP"))

        if endereco:
            elementos.append(f"residente e {domiciliado} na(o) {formatar_nome_proprio(endereco)}")
        if bairro:
            elementos.append(f"Bairro {formatar_nome_proprio(bairro)}")

        if cidade and uf:
            elementos.append(f"{formatar_nome_proprio(cidade)}/{uf}")
        elif cidade:
            elementos.append(formatar_nome_proprio(cidade))
        elif uf:
            elementos.append(uf)

        if cep:
            elementos.append(f"CEP {cep}")

        textos.append((nome, limpar_espacos(", ".join(elementos))))
    return textos

# =========================
# DOCX helpers
# =========================
def add_paragrafo(doc: Document, texto: str, *, negrito=False, maiusculo=False, tamanho=12,
                  centralizado=False, titulo=False):
    texto = limpar_espacos(texto)
    if maiusculo:
        texto = texto.upper()
    par = doc.add_paragraph()
    run = par.add_run(texto)

    run.font.name = "Book Antiqua"
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Book Antiqua")
    except Exception:
        pass
    run.font.size = Pt(tamanho)
    run.bold = negrito

    par.alignment = WD_ALIGN_PARAGRAPH.CENTER if centralizado else WD_ALIGN_PARAGRAPH.JUSTIFY
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.first_line_indent = Cm(0 if titulo else 1.25)
    return par

def add_titulo_numerado(doc: Document, indice: int, titulo: str):
    add_paragrafo(doc, "", titulo=True)
    add_paragrafo(doc, f"{indice} - {titulo}", negrito=True, maiusculo=True, tamanho=14, titulo=True)
    add_paragrafo(doc, "", titulo=True)

def run_book(paragraph, text, *, bold=False, size=12):
    r = paragraph.add_run(text)
    r.bold = bold
    r.font.name = "Book Antiqua"
    try:
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Book Antiqua")
    except Exception:
        pass
    r.font.size = Pt(size)
    return r

# =========================
# PRINCIPAL
# =========================
def main():
    partes_df, processo_df, advogado_df, secoes_opcionais_df = ler_planilhas(ARQUIVO_EXCEL)

    processo = processo_df.iloc[0] if not processo_df.empty else pd.Series(dtype=object)
    comarca_completa = valor_str(processo.get("Comarca"))
    comarca = re.sub(r"\s*[-–]?\s*paraíba\s*$", "", comarca_completa, flags=re.IGNORECASE).strip() or comarca_completa
    vara = valor_str(processo.get("Vara")).upper()
    tipo_causa = valor_str(processo.get("Tipo de Causa"))
    fatos = valor_str(processo.get("Fatos"))
    valor_causa = processo.get("Valor da Causa", 0)

    adv = advogado_df.iloc[0] if not advogado_df.empty else pd.Series(dtype=object)
    nome_adv = valor_str(adv.get("Nome Completo")).upper()
    uf_adv   = valor_str(adv.get("UF")).upper()
    oab_adv  = valor_str(adv.get("OAB"))

    def tipo_norm(x): return valor_str(x).lower().strip()
    requerentes = partes_df[partes_df.get("Tipo", "").apply(tipo_norm) == "requerente"].copy()
    requeridos  = partes_df[partes_df.get("Tipo", "").apply(tipo_norm) == "requerido"].copy()

    qual_reqs = gerar_qualificacao(requerentes)
    qual_reus = gerar_qualificacao(requeridos)

    try:
        documento = Document(ARQUIVO_MODELO)
    except Exception as e:
        raise RuntimeError(f"Erro ao abrir o modelo '{ARQUIVO_MODELO}': {e}")

    # Endereçamento
    add_paragrafo(
        documento,
        f"Excelentíssima Senhora Doutora Juíza de Direito da {vara} da Comarca de {comarca} - Paraíba",
        negrito=True, maiusculo=True, tamanho=16, centralizado=True, titulo=True
    )
    add_paragrafo(documento, ""); add_paragrafo(documento, ""); add_paragrafo(documento, "")

    tipo_acao_formatado = tipo_causa.upper()
    if not tipo_acao_formatado.startswith("AÇÃO DE"):
        tipo_acao_formatado = f"AÇÃO DE {tipo_acao_formatado}"

    par = documento.add_paragraph()
    par.paragraph_format.line_spacing = 1.5
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Requerentes + representação + ação
    for i, (nome, qual) in enumerate(qual_reqs):
        run_book(par, nome, bold=True)
        run_book(par, ", " + qual)
        if i < len(qual_reqs) - 1:
            run_book(par, "; ")
        else:
            run_book(par, "; neste ato representado por ")
            run_book(par, nome_adv, bold=True)
            run_book(par, f", inscrito na OAB/{uf_adv} sob o nº {oab_adv}, vem, respeitosamente, à presença de Vossa Excelência propor a presente ")
            run_book(par, tipo_acao_formatado, bold=True)

    if qual_reus:
        req_texto = ", ".join([f"{nome}, {qual}" for nome, qual in qual_reus])
        run_book(par, f", em face de {req_texto}, pelos fatos e fundamentos que passa a expor:")
    else:
        run_book(par, ", pelos fatos e fundamentos que passa a expor:")

    # Seções
    indice = 1
    textos_para_extracao_dispositivos = []

    # 1 - DOS FATOS (GPT com fallback para 'fatos')
    prompt_fatos = (
        f"Você é um advogado especialista. Reescreva os fatos da ação de {tipo_causa} de forma clara, objetiva e formal:\n\n{fatos}"
    )
    fatos_interpretados = gpt_or_fallback(prompt_fatos, fallback=fatos, max_tokens=1500, temperature=0.2)
    add_titulo_numerado(documento, indice, "Dos Fatos"); indice += 1
    for linha in fatos_interpretados.splitlines():
        if linha.strip():
            add_paragrafo(documento, linha.strip())
    textos_para_extracao_dispositivos.append(fatos_interpretados)

    # 2 - GRATUIDADE (antirrepetição)
    if INCLUIR_GRATUIDADE:
        def _prompt_grat_fn(fatos_ctx: str, instr: str) -> str:
            return montar_prompt_gratuidade(fatos_ctx, instr)

        texto_gratuidade = gerar_texto_unico(
            secao_key="gratuidade",
            construir_prompt_fn=_prompt_grat_fn,
            fatos_contexto=fatos_interpretados or fatos,
            tentativas=6,
            alvo_max_similaridade=0.33,
            ngram_n=3,
            temperaturas=(0.4, 0.6, 0.8, 0.9, 1.0, 1.1),
            variantes_instrucao=VARIANTES_GRATUIDADE,
            fallback_padrao=(
                "À luz do art. 5º, LXXIV, da Constituição Federal e dos arts. 98 a 102 do CPC, requer-se a concessão "
                "dos benefícios da justiça gratuita, dado que a parte não possui condições de suportar as despesas "
                "processuais sem prejuízo de seu sustento, bastando, para tanto, a declaração de hipossuficiência, "
                "salvo prova em contrário."
            )
        )

        add_titulo_numerado(documento, indice, "Da Gratuidade da Justiça"); indice += 1
        for ln in texto_gratuidade.splitlines():
            if ln.strip():
                add_paragrafo(documento, ln.strip())
        textos_para_extracao_dispositivos.append(texto_gratuidade)

    # 3 - FUNDAMENTOS (antirrepetição)
    def _prompt_fund_fn(fatos_ctx: str, instr: str) -> str:
        return montar_prompt_fundamentos(fatos_ctx, instr, tipo_causa)

    fundamentos = gerar_texto_unico(
        secao_key="fundamentos",
        construir_prompt_fn=_prompt_fund_fn,
        fatos_contexto=fatos_interpretados or fatos,
        tentativas=6,
        alvo_max_similaridade=0.34,
        ngram_n=3,
        temperaturas=(0.35, 0.55, 0.75, 0.9, 1.0, 1.05),
        variantes_instrucao=VARIANTES_FUNDAMENTOS,
        fallback_padrao=(
            "A controvérsia deve ser solucionada à luz da Constituição Federal e do Código de Processo Civil, "
            "respeitando-se a inafastabilidade da jurisdição (art. 5º, XXXV, CF), o devido processo legal e o "
            "contraditório substancial (art. 5º, LIV e LV, CF), bem como a cooperação e a efetividade "
            "da tutela jurisdicional previstas no CPC."
        )
    )

    add_titulo_numerado(documento, indice, "Dos Fundamentos Jurídicos"); indice += 1
    for ln in fundamentos.splitlines():
        if ln.strip():
            add_paragrafo(documento, ln.strip())
    textos_para_extracao_dispositivos.append(fundamentos)

    # 4+ - SEÇÕES OPCIONAIS
    if INCLUIR_SECOES_OPCIONAIS and not secoes_opcionais_df.empty:
        for _, row in secoes_opcionais_df.iterrows():
            titulo = valor_str(row.get("Título"))
            observ  = valor_str(row.get("Observação"))
            if not titulo:
                continue
            prompt_secao = (
                f"Você é um advogado especialista. Redija a seção '{titulo}', com base na legislação brasileira e nas "
                f"observações: \"{observ}\". A ação trata de {tipo_causa}. Linguagem formal e estruturada. "
                f"Responda só com os parágrafos (sem repetir o título)."
            )
            conteudo_secao = gpt_or_fallback(prompt_secao, fallback=observ or "")
            add_titulo_numerado(documento, indice, titulo); indice += 1
            for ln in conteudo_secao.splitlines():
                if ln.strip() and not ln.strip().upper().startswith(titulo.upper()):
                    add_paragrafo(documento, ln.strip())
            textos_para_extracao_dispositivos.append(conteudo_secao)

    # PEDIDOS (antirrepetição)
    def _prompt_ped_fn(fatos_ctx: str, instr: str) -> str:
        return montar_prompt_pedidos(fatos_ctx, instr, tipo_causa)

    pedidos_brutos = gerar_texto_unico(
        secao_key="pedidos",
        construir_prompt_fn=_prompt_ped_fn,
        fatos_contexto=fatos_interpretados or fatos,
        tentativas=7,
        alvo_max_similaridade=0.32,
        ngram_n=3,
        temperaturas=(0.35, 0.55, 0.75, 0.9, 1.0, 1.05, 1.1),
        variantes_instrucao=VARIANTES_PEDIDOS,
        fallback_padrao=(
            "- A citação do(s) requerido(s) para responder, sob pena de revelia;\n"
            "- A procedência dos pedidos, conforme fundamentação;\n"
            "- A condenação do(s) requerido(s) ao pagamento de custas e honorários (art. 85, CPC)."
        )
    )

    add_titulo_numerado(documento, indice, "Dos Pedidos"); indice += 1
    add_paragrafo(documento, "Diante do exposto, requerem a Vossa Excelência:", titulo=True)

    itens = [re.sub(r"^\(?[a-zA-Z0-9]+\)?[.)]?\s*", "", ln.strip()) for ln in pedidos_brutos.splitlines() if ln.strip()]
    for i, item in enumerate(itens):
        letra = chr(97 + i)  # a, b, c...
        add_paragrafo(documento, f"{letra}) {item}")

    # Valor da causa + fechamento
    add_paragrafo(documento, "")
    add_paragrafo(documento, f"Dá-se à presente causa o valor de R$ {moeda_brasil(valor_causa)}", titulo=True)
    add_paragrafo(documento, "")
    add_paragrafo(documento, "Nesses termos,", titulo=True)
    add_paragrafo(documento, "Pede deferimento.", titulo=True)
    add_paragrafo(documento, "")

    # Assinatura
    add_paragrafo(documento, formatar_nome_proprio(nome_adv), tamanho=12, negrito=True, titulo=True)
    add_paragrafo(documento, f"OAB/{uf_adv} {oab_adv}", tamanho=12, negrito=True, titulo=True)

    # Opcional: dispositivos citados
    if INCLUIR_SUMARIO_DISPOSITIVOS:
        corpo_total = "\n".join(textos_para_extracao_dispositivos)
        dispositivos = extrair_artigos_de_lei(corpo_total)
        if dispositivos:
            add_paragrafo(documento, "")
            add_titulo_numerado(documento, indice, "Dispositivos Citados"); indice += 1
            add_paragrafo(documento, "Para pronta referência, destacam-se os dispositivos mencionados:")
            for d in dispositivos:
                add_paragrafo(documento, f"• {d}")

    # Salvar com retentativas
    for tentativa in range(1, TENTATIVAS_SALVAR + 1):
        try:
            documento.save(ARQUIVO_SAIDA)
            print(f"✅ Petição gerada com sucesso em: {ARQUIVO_SAIDA}")
            break
        except PermissionError:
            print(f"❌ Arquivo aberto? Feche-o. Tentando novamente ({tentativa}/{TENTATIVAS_SALVAR})...")
            time.sleep(TEMPO_ESPERA_SALVAR_SEG)
    else:
        print("❌ Não foi possível salvar. Verifique se o arquivo está fechado e o caminho é válido.")

if __name__ == "__main__":
    main()
